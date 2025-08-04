import streamlit as st
import pandas as pd
from redminelib import Redmine
import requests
import datetime
import os
import re
from docx import Document
from docx.shared import Inches
import io

# Config dosyasını import et
from config import (
    REDMINE_API_KEY,
    REDMINE_API_URL,
    MATTERMOST_TOKEN,
    MATTERMOST_BASE_URL,
    TARGET_MATTERMOST_CHANNELS
)

# ==============================================================================
# 1. REDMINE VERİ ÇEKME FONKSİYONU (OPTİMİZE EDİLMİŞ)
# ==============================================================================
def fast_redmine_fetch(search_term):
    """
    Hızlı Redmine arama - sadece başlık ve açıklamada arama yapar
    """
    print(f"Connecting to Redmine ({REDMINE_API_URL})...")
    collected_data = []
    
    try:
        redmine = Redmine(REDMINE_API_URL, key=REDMINE_API_KEY)
        print("Login successful. Fetching all issues for fast search...")

        # Sadece temel verileri çek (journals ve attachments olmadan)
        issues = redmine.issue.filter(status_id='*')
        
        print(f"Total issues fetched: {len(issues)}. Now filtering for '{search_term}'...")
        
        for issue in issues:
            subject = getattr(issue, 'subject', '') or ''
            description = getattr(issue, 'description', '') or ''
            
            is_match = search_term.lower() in subject.lower() or search_term.lower() in description.lower()
            
            if is_match:
                tracker_name = getattr(issue, 'tracker', {}).name
                status_name = getattr(issue, 'status', {}).name
                author_name = getattr(issue, 'author', {}).name

                collected_data.append({
                    "Source_Platform": "Redmine", 
                    "ID": issue.id,
                    "Content_Type": f"Tracker: {tracker_name} (Status: {status_name})",
                    "Title": subject, 
                    "Description": description,
                    "Author": author_name, 
                    "Creation_Date": issue.created_on.strftime('%Y-%m-%d %H:%M:%S'),
                    "Notes": "N/A", 
                    "Attached_Files": "N/A",
                    "Channel_ID": "N/A"
                })
            
    except Exception as e:
        print(f"!!!! REDMINE ERROR: {e} !!!!")
        
    return collected_data

# ==============================================================================
# 2. MATTERMOST VERİ ÇEKME FONKSİYONU (OPTİMİZE EDİLMİŞ)
# ==============================================================================
def fast_mattermost_fetch(search_term):
    print(f"Connecting to Mattermost and searching in {len(TARGET_MATTERMOST_CHANNELS)} specific channels...")
    base_url = MATTERMOST_BASE_URL
    headers = {"Authorization": f"Bearer {MATTERMOST_TOKEN}"}
    collected_data = []
    processed_thread_ids = set()
    
    # Sadece # ile başlayan aramaları kabul et
    if not search_term.startswith('#'):
        print(f"Search term '{search_term}' does not start with '#'. Only hashtag searches are allowed.")
        return collected_data
    
    try:
        teams_response = requests.get(f"{base_url}/users/me/teams", headers=headers)
        teams_response.raise_for_status()
        teams = teams_response.json()
        if not teams: return []
        team_id = teams[0]['id']
        
        # Önce normal arama yap (Mattermost API search)
        search_url = f"{base_url}/teams/{team_id}/posts/search"
        payload = {"terms": search_term, "is_or_search": True}
        search_response = requests.post(search_url, headers=headers, json=payload)
        search_response.raise_for_status()
        search_results = search_response.json()
        
        print(f"Found {len(search_results.get('posts', {}))} potential posts with '{search_term}'. Processing threads...")
        
        # Bulunan thread'leri işle
        for post in search_results.get('posts', {}).values():
            if post['channel_id'] not in TARGET_MATTERMOST_CHANNELS:
                continue

            thread_id = post.get('root_id') or post['id']
            if thread_id in processed_thread_ids: continue
            processed_thread_ids.add(thread_id)
            
            try:
                # Tüm thread'i çek
                thread_response = requests.get(f"{base_url}/posts/{thread_id}/thread", headers=headers)
                thread_response.raise_for_status()
                full_thread = thread_response.json()
                
                ordered_post_ids = full_thread.get('order', [])
                if not ordered_post_ids: continue
                
                relevant_messages = []
                root_post = full_thread['posts'][ordered_post_ids[0]]
                root_date = datetime.datetime.fromtimestamp(root_post['create_at'] / 1000).strftime('%Y-%m-%d %H:%M:%S')
                root_author = root_post.get('user_id', 'User')
                
                # Kök mesajı her zaman ekle (bağlam için)
                relevant_messages.append(f"--- KONU BAŞLANGICI ---\n[{root_date} - User {root_author}]:\n{root_post.get('message', '')}")
                
                # Thread içindeki tüm mesajları kontrol et
                thread_contains_killed_prey = False
                
                # Kök mesajda "killed a prey" var mı?
                if "killed a prey" in root_post.get('message', '').lower() or "killed a prey:" in root_post.get('message', '').lower():
                    thread_contains_killed_prey = True
                
                # Thread içindeki tüm mesajları kontrol et
                for post_id in ordered_post_ids[1:]:
                    reply = full_thread['posts'][post_id]
                    reply_message = reply.get('message', '')
                    
                    # Arama terimi bu mesajda var mı? (bağlam için ekle)
                    if search_term.lower() in reply_message.lower():
                        reply_date = datetime.datetime.fromtimestamp(reply['create_at'] / 1000).strftime('%H:%M:%S')
                        reply_author = reply.get('user_id', 'User')
                        relevant_messages.append(f"--- İLGİLİ YANIT ---\n[{reply_date} - User {reply_author}]:\n{reply_message}")
                    
                    # "killed a prey" bu mesajda var mı?
                    if "killed a prey" in reply_message.lower() or "killed a prey:" in reply_message.lower():
                        thread_contains_killed_prey = True
                        reply_date = datetime.datetime.fromtimestamp(reply['create_at'] / 1000).strftime('%H:%M:%S')
                        reply_author = reply.get('user_id', 'User')
                        relevant_messages.append(f"--- İLGİLİ YANIT (KILLED A PREY) ---\n[{reply_date} - User {reply_author}]:\n{reply_message}")
                
                # Thread'i ekle (killed a prey durumuna göre)
                if len(relevant_messages) > 1:  # En az kök mesajı ve bir yanıt var
                    # Thread'in durumunu belirle
                    status = "tamamlandi" if thread_contains_killed_prey else "devam_ediyor"
                    
                    collected_data.append({
                        "Source_Platform": "Mattermost", 
                        "ID": root_post.get('id', 'Unknown'),
                        "Title": f"Mattermost Konusu: {root_post.get('id', 'Unknown')}",
                        "Description": "\n\n".join(relevant_messages),
                        "Author": root_author, 
                        "Creation_Date": root_date, 
                        "Content_Type": "Mattermost Thread",
                        "Channel_ID": root_post.get('channel_id', 'Unknown'),
                        "Notes": "N/A", 
                        "Attached_Files": "N/A",
                        "Status": status  # Durum bilgisini ekle
                    })
                
            except Exception as thread_error:
                print(f"Warning: Could not process Mattermost thread for ID {thread_id}. Error: {thread_error}")
                
    except requests.exceptions.RequestException as e:
        print(f"!!!! MATTERMOST ERROR: {e} !!!!")
        
    return collected_data

# ==============================================================================
# 3. ANA VERİ YÖNETİM FONKSİYONU (OPTİMİZE EDİLMİŞ)
# ==============================================================================
@st.cache_data(ttl=1800)  # 30 dakika önbellek
def get_all_data_fast(search_term, platform_filter="Tümü"):
    """
    Hızlı veri çekme - platform filtresine göre optimize edilmiş versiyon
    """
    print(f"'{search_term}' için hızlı veri çekme başlatılıyor... (Platform: {platform_filter})")
    
    # Platform filtresine göre veri çek
    if platform_filter == "Redmine":
        print("Sadece Redmine verileri çekiliyor...")
        redmine_list = fast_redmine_fetch(search_term)
        mattermost_list = []
    elif platform_filter == "Mattermost":
        print("Sadece Mattermost verileri çekiliyor...")
        redmine_list = []
        mattermost_list = fast_mattermost_fetch(search_term)
    else:  # "Tümü" seçiliyse
        print("Tüm platformlardan veri çekiliyor...")
        redmine_list = fast_redmine_fetch(search_term)
        mattermost_list = fast_mattermost_fetch(search_term)
    
    redmine_df = pd.DataFrame(redmine_list)
    mattermost_df = pd.DataFrame(mattermost_list)
    
    combined_df = pd.concat([redmine_df, mattermost_df], ignore_index=True)

    print(f"Veri çekme tamamlandı. Toplam {len(combined_df)} sonuç bulundu.")
    return combined_df

# ==============================================================================
# 4. CHATBOT FONKSİYONLARI
# ==============================================================================
def parse_natural_language_query(query):
    """
    Doğal dil sorgusunu parse eder ve arama parametrelerini çıkarır
    """
    query = query.lower()
    
    # Varsayılan değerler
    search_term = None
    platform_filter = "Tümü"
    status_filter = "Tümü"
    date_filter = None
    
    # Hashtag arama
    hashtag_pattern = r'#(\w+)'
    hashtags = re.findall(hashtag_pattern, query)
    if hashtags:
        search_term = f"#{hashtags[0]}"  # İlk hashtag'i al
    
    # Platform filtreleme
    if 'redmine' in query:
        platform_filter = "Redmine"
    elif 'mattermost' in query:
        platform_filter = "Mattermost"
    
    # Durum filtreleme
    if any(word in query for word in ['tamamlanan', 'tamamlandı', 'completed', 'finished']):
        status_filter = "Tamamlanan"
    elif any(word in query for word in ['devam eden', 'ongoing', 'continuing']):
        status_filter = "Devam Eden"
    
    # Tarih filtreleme (basit)
    if 'bugün' in query or 'today' in query:
        date_filter = datetime.date.today()
    elif 'dün' in query or 'yesterday' in query:
        date_filter = datetime.date.today() - datetime.timedelta(days=1)
    elif 'bu hafta' in query or 'this week' in query:
        date_filter = datetime.date.today() - datetime.timedelta(days=7)
    
    return {
        'search_term': search_term,
        'platform_filter': platform_filter,
        'status_filter': status_filter,
        'date_filter': date_filter,
        'original_query': query
    }

def create_word_document(data, query_info):
    """
    Arama sonuçlarını Word dokümanına aktarır
    """
    doc = Document()
    
    # Başlık
    title = doc.add_heading('İş Takip Arama Raporu', 0)
    title.alignment = 1  # Ortalanmış
    
    # Sorgu bilgileri
    doc.add_heading('Sorgu Bilgileri', level=1)
    p = doc.add_paragraph()
    p.add_run('Orijinal Sorgu: ').bold = True
    p.add_run(query_info['original_query'])
    
    p = doc.add_paragraph()
    p.add_run('Arama Terimi: ').bold = True
    p.add_run(query_info['search_term'] or 'Belirtilmemiş')
    
    p = doc.add_paragraph()
    p.add_run('Platform Filtresi: ').bold = True
    p.add_run(query_info['platform_filter'])
    
    p = doc.add_paragraph()
    p.add_run('Durum Filtresi: ').bold = True
    p.add_run(query_info['status_filter'])
    
    if query_info['date_filter']:
        p = doc.add_paragraph()
        p.add_run('Tarih Filtresi: ').bold = True
        p.add_run(str(query_info['date_filter']))
    
    # İstatistikler
    doc.add_heading('İstatistikler', level=1)
    p = doc.add_paragraph()
    p.add_run(f'Toplam Sonuç: {len(data)}').bold = True
    
    redmine_count = len(data[data['Source_Platform'] == 'Redmine'])
    mattermost_count = len(data[data['Source_Platform'] == 'Mattermost'])
    
    p = doc.add_paragraph()
    p.add_run(f'Redmine Sonuçları: {redmine_count}')
    
    p = doc.add_paragraph()
    p.add_run(f'Mattermost Sonuçları: {mattermost_count}')
    
    # Mattermost durum dağılımı
    if mattermost_count > 0:
        mattermost_data = data[data['Source_Platform'] == 'Mattermost']
        tamamlanan_count = len(mattermost_data[mattermost_data['Status'] == 'tamamlandi'])
        devam_eden_count = len(mattermost_data[mattermost_data['Status'] == 'devam_ediyor'])
        
        p = doc.add_paragraph()
        p.add_run('Mattermost Durum Dağılımı:').bold = True
        
        p = doc.add_paragraph()
        p.add_run(f'• Tamamlanan: {tamamlanan_count}')
        
        p = doc.add_paragraph()
        p.add_run(f'• Devam Eden: {devam_eden_count}')
    
    # Detaylı Sonuçlar
    doc.add_heading('Detaylı Sonuçlar', level=1)
    
    # Redmine sonuçları
    redmine_data = data[data['Source_Platform'] == 'Redmine']
    if not redmine_data.empty:
        doc.add_heading('Redmine Sonuçları', level=2)
        for index, row in redmine_data.iterrows():
            doc.add_heading(f"🔴 {row.get('Title', 'Başlık Yok')}", level=3)
            
            p = doc.add_paragraph()
            p.add_run('📅 Tarih: ').bold = True
            p.add_run(str(row.get('Creation_Date', 'N/A')))
            
            p = doc.add_paragraph()
            p.add_run('👤 Yazar: ').bold = True
            p.add_run(str(row.get('Author', 'N/A')))
            
            p = doc.add_paragraph()
            p.add_run('🏷️ Tip: ').bold = True
            p.add_run(str(row.get('Content_Type', 'N/A')))
            
            p = doc.add_paragraph()
            p.add_run('📝 Açıklama:').bold = True
            p.add_run('\n' + str(row.get('Description', '')))
            
            doc.add_paragraph()  # Boşluk
    
    # Mattermost sonuçları
    mattermost_data = data[data['Source_Platform'] == 'Mattermost']
    if not mattermost_data.empty:
        doc.add_heading('Mattermost Sonuçları', level=2)
        
        # Tamamlanan işler
        tamamlanan_data = mattermost_data[mattermost_data['Status'] == 'tamamlandi']
        if not tamamlanan_data.empty:
            doc.add_heading('🟢 Tamamlanan İşler', level=3)
            for index, row in tamamlanan_data.iterrows():
                doc.add_heading(f"✅ {row.get('Title', 'Başlık Yok')}", level=4)
                
                p = doc.add_paragraph()
                p.add_run('📅 Tarih: ').bold = True
                p.add_run(str(row.get('Creation_Date', 'N/A')))
                
                p = doc.add_paragraph()
                p.add_run('👤 Yazar: ').bold = True
                p.add_run(str(row.get('Author', 'N/A')))
                
                p = doc.add_paragraph()
                p.add_run('📝 Mesajlar:').bold = True
                p.add_run('\n' + str(row.get('Description', '')))
                
                doc.add_paragraph()  # Boşluk
        
        # Devam eden işler
        devam_eden_data = mattermost_data[mattermost_data['Status'] == 'devam_ediyor']
        if not devam_eden_data.empty:
            doc.add_heading('🔵 Devam Eden İşler', level=3)
            for index, row in devam_eden_data.iterrows():
                doc.add_heading(f"⏳ {row.get('Title', 'Başlık Yok')}", level=4)
                
                p = doc.add_paragraph()
                p.add_run('📅 Tarih: ').bold = True
                p.add_run(str(row.get('Creation_Date', 'N/A')))
                
                p = doc.add_paragraph()
                p.add_run('👤 Yazar: ').bold = True
                p.add_run(str(row.get('Author', 'N/A')))
                
                p = doc.add_paragraph()
                p.add_run('📝 Mesajlar:').bold = True
                p.add_run('\n' + str(row.get('Description', '')))
                
                doc.add_paragraph()  # Boşluk
    
    # Rapor oluşturma tarihi
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Rapor Oluşturma Tarihi: ').bold = True
    p.add_run(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    return doc

# ==============================================================================
# 5. STREAMLIT ARAYÜZÜ
# ==============================================================================
def main():
    st.set_page_config(layout="wide", page_title="İş Takip Arama Aracı")
    
    st.title('🔍 İş Takip Arama Aracı')
    st.markdown("Redmine ve Mattermost platformlarında iş isimlerine göre arama yapın.")
    
    # Session state ile veri saklama
    if 'search_results' not in st.session_state:
        st.session_state.search_results = None
    if 'search_term' not in st.session_state:
        st.session_state.search_term = ""
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'last_results' not in st.session_state:
        st.session_state.last_results = None
    
    # Sekmeler oluştur
    tab1, tab2 = st.tabs(["🔍 Manuel Arama", "🤖 Chatbot Arama"])
    
    # Manuel Arama Sekmesi
    with tab1:
        manual_search_interface()
    
    # Chatbot Arama Sekmesi
    with tab2:
        chatbot_search_interface()
    
def manual_search_interface():
    """Manuel arama arayüzü"""
    # Ana arama bölümü
    st.header("📝 Arama Yap")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        search_term = st.text_input(
            "Aranacak iş ismini girin:",
            placeholder="Örnek: ATP, pharmacircle, yama notu...",
            value=st.session_state.search_term
        )
    
    with col2:
        st.write("")
        st.write("")
        search_button = st.button("🔍 Ara", type="primary")
    
    # Filtreleme seçenekleri - her zaman görünür
    st.header("🔧 Filtreleme Seçenekleri")
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        platform_filter = st.selectbox(
            "Platform Filtresi:",
            ["Tümü", "Redmine", "Mattermost"]
        )
    
    with col2:
        text_filter = st.text_input("İçerik içinde ara:")
    
    with col3:
        # Mattermost durum filtresi
        st.write("📊 Mattermost Durumu:")
        mattermost_status_filter = st.selectbox(
            "Durum Filtresi:",
            ["Tümü", "Tamamlanan", "Devam Eden"],
            key="mattermost_status"
        )
    
    with col4:
        # Tarih filtresi
        st.write("📅 Tarih Aralığı:")
        date_filter_enabled = st.checkbox("Tarih filtresi kullan", value=False)
        
        if date_filter_enabled:
            col_date1, col_date2 = st.columns(2)
            with col_date1:
                start_date = st.date_input("Başlangıç tarihi:", value=None)
            with col_date2:
                end_date = st.date_input("Bitiş tarihi:", value=None)
        else:
            start_date = None
            end_date = None
    
    with col5:
        st.write("")
        st.write("")
        clear_cache = st.button("🔄 Önbelleği Temizle")
        if clear_cache:
            st.cache_data.clear()
            st.session_state.search_results = None
            st.rerun()
    
    # Arama yapma
    if search_button and search_term:
        st.session_state.search_term = search_term
        
        with st.spinner(f"'{search_term}' için veriler aranıyor... (Bu işlem 1-2 dakika sürebilir)"):
            try:
                # Progress bar ekle
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Platform filtresine göre mesaj göster
                if platform_filter == "Redmine":
                    status_text.text("Redmine verileri çekiliyor...")
                elif platform_filter == "Mattermost":
                    status_text.text("Mattermost verileri çekiliyor...")
                else:
                    status_text.text("Redmine verileri çekiliyor...")
                
                progress_bar.progress(25)
                
                df = get_all_data_fast(search_term, platform_filter)
                st.session_state.search_results = df
                
                progress_bar.progress(100)
                status_text.text("Arama tamamlandı!")
                
                if df.empty:
                    st.warning(f"'{search_term}' için hiçbir sonuç bulunamadı.")
                else:
                    st.success(f"'{search_term}' için {len(df)} sonuç bulundu!")
                    
                    # Filtreleme uygula (platform filtresi zaten uygulandı)
                    filtered_df = df.copy()
                    
                    if text_filter:
                        mask = (filtered_df['Title'].str.contains(text_filter, case=False, na=False) | 
                               filtered_df['Description'].str.contains(text_filter, case=False, na=False))
                        filtered_df = filtered_df[mask]
                    
                    # Mattermost durum filtresi uygula
                    if mattermost_status_filter != "Tümü":
                        filtered_df = filter_mattermost_by_status(filtered_df, mattermost_status_filter)
                    
                    # Tarih filtresi uygula
                    if date_filter_enabled and (start_date is not None or end_date is not None):
                        # Creation_Date sütununu datetime'a çevir
                        filtered_df['Creation_Date_dt'] = pd.to_datetime(filtered_df['Creation_Date'], errors='coerce')
                        
                        if start_date is not None:
                            start_date_dt = pd.to_datetime(start_date)
                            filtered_df = filtered_df[filtered_df['Creation_Date_dt'] >= start_date_dt]
                        
                        if end_date is not None:
                            end_date_dt = pd.to_datetime(end_date) + pd.Timedelta(days=1)  # Bitiş gününü dahil et
                            filtered_df = filtered_df[filtered_df['Creation_Date_dt'] < end_date_dt]
                        
                        # Geçici sütunu sil
                        filtered_df = filtered_df.drop('Creation_Date_dt', axis=1)
                    
                    # Sonuçları göster
                    st.header("📊 Bulunan Sonuçlar")
                    st.subheader(f"Filtrelenmiş Sonuçlar ({len(filtered_df)} adet)")
                    
                    if not filtered_df.empty:
                        # Sonuçları tarihe göre sırala
                        filtered_df = filtered_df.sort_values(by='Creation_Date', ascending=False)
                        
                        # Mattermost sonuçları için durum analizi yap
                        if len(filtered_df[filtered_df['Source_Platform'] == 'Mattermost']) > 0:
                            mattermost_df = filtered_df[filtered_df['Source_Platform'] == 'Mattermost'].copy()
                            
                            # Durumlara göre grupla (Status sütunu zaten fast_mattermost_fetch'te set edildi)
                            tamamlanan = mattermost_df[mattermost_df['Status'] == 'tamamlandi']
                            devam_eden = mattermost_df[mattermost_df['Status'] == 'devam_ediyor']
                            
                            # Redmine sonuçları
                            redmine_df = filtered_df[filtered_df['Source_Platform'] == 'Redmine']
                            
                            # Sekmeler oluştur
                            tab1, tab2, tab3 = st.tabs([
                                f"🔴 Redmine ({len(redmine_df)})",
                                f"🟢 Tamamlanan ({len(tamamlanan)})",
                                f"🔵 Devam Eden ({len(devam_eden)})"
                            ])
                            
                            # Redmine sekmesi
                            with tab1:
                                if not redmine_df.empty:
                                    for index, row in redmine_df.iterrows():
                                        with st.container():
                                            st.markdown(f"### 🔴 Redmine - {row.get('Title', 'Başlık Yok')}")
                                            with st.expander("📋 Detayları Gör", expanded=False):
                                                col1, col2 = st.columns([2, 1])
                                                with col1:
                                                    st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                                                    st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                                                    st.markdown(f"**🏷️ Tip:** {row.get('Content_Type', 'N/A')}")
                                                    st.markdown(f"**🆔 ID:** {row.get('ID', 'N/A')}")
                                                with col2:
                                                    st.markdown("**📊 İstatistikler:**")
                                                    description_length = len(str(row.get('Description', '')))
                                                    st.markdown(f"• Açıklama uzunluğu: {description_length} karakter")
                                                st.markdown("---")
                                                st.markdown("**📝 Açıklama:**")
                                                st.text_area("", value=row.get('Description', ''), height=200, disabled=True)
                                            st.markdown("---")
                                else:
                                    st.info("Redmine sonucu bulunamadı.")
                            
                            # Tamamlanan sekmesi (Yeşil)
                            with tab2:
                                if not tamamlanan.empty:
                                    for index, row in tamamlanan.iterrows():
                                        with st.container():
                                            st.markdown(f"### 🟢 Tamamlanan - {row.get('Title', 'Başlık Yok')}")
                                            with st.expander("📋 Detayları Gör", expanded=False):
                                                col1, col2 = st.columns([2, 1])
                                                with col1:
                                                    st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                                                    st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                                                    st.markdown(f"**📺 Kanal ID:** {row.get('Channel_ID', 'N/A')}")
                                                with col2:
                                                    st.markdown("**📊 İstatistikler:**")
                                                    description_length = len(str(row.get('Description', '')))
                                                    st.markdown(f"• Açıklama uzunluğu: {description_length} karakter")
                                                st.markdown("---")
                                                st.markdown("**📝 Mesajlar:**")
                                                st.text_area("", value=row.get('Description', ''), height=200, disabled=True)
                                            st.markdown("---")
                                else:
                                    st.info("Tamamlanan iş bulunamadı.")
                            
                            # Devam Eden sekmesi (Mavi)
                            with tab3:
                                if not devam_eden.empty:
                                    for index, row in devam_eden.iterrows():
                                        with st.container():
                                            st.markdown(f"### 🔵 Devam Eden - {row.get('Title', 'Başlık Yok')}")
                                            with st.expander("📋 Detayları Gör", expanded=False):
                                                col1, col2 = st.columns([2, 1])
                                                with col1:
                                                    st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                                                    st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                                                    st.markdown(f"**📺 Kanal ID:** {row.get('Channel_ID', 'N/A')}")
                                                with col2:
                                                    st.markdown("**📊 İstatistikler:**")
                                                    description_length = len(str(row.get('Description', '')))
                                                    st.markdown(f"• Açıklama uzunluğu: {description_length} karakter")
                                                st.markdown("---")
                                                st.markdown("**📝 Mesajlar:**")
                                                st.text_area("", value=row.get('Description', ''), height=200, disabled=True)
                                            st.markdown("---")
                                else:
                                    st.info("Devam eden iş bulunamadı.")
                        else:
                            # Sadece Redmine sonuçları varsa normal gösterim
                            for index, row in filtered_df.iterrows():
                                with st.container():
                                    platform_badge = "🔴 Redmine" if row['Source_Platform'] == 'Redmine' else "💬 Mattermost"
                                    st.markdown(f"### {platform_badge} - {row.get('Title', 'Başlık Yok')}")
                                    with st.expander("📋 Detayları Gör", expanded=False):
                                        col1, col2 = st.columns([2, 1])
                                        with col1:
                                            st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                                            st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                                            st.markdown(f"**🏷️ Tip:** {row.get('Content_Type', 'N/A')}")
                                            if row.get('Source_Platform') == 'Redmine':
                                                st.markdown(f"**🆔 ID:** {row.get('ID', 'N/A')}")
                                            else:
                                                st.markdown(f"**📺 Kanal ID:** {row.get('Channel_ID', 'N/A')}")
                                        with col2:
                                            st.markdown("**📊 İstatistikler:**")
                                            description_length = len(str(row.get('Description', '')))
                                            st.markdown(f"• Açıklama uzunluğu: {description_length} karakter")
                                        st.markdown("---")
                                        st.markdown("**📝 Açıklama/Mesajlar:**")
                                        st.text_area("", value=row.get('Description', ''), height=200, disabled=True)
                                    st.markdown("---")
                    
                    # İstatistikler
                    st.header("📈 İstatistikler")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Toplam Sonuç", len(df))
                    
                    with col2:
                        redmine_count = len(df[df['Source_Platform'] == 'Redmine'])
                        st.metric("Redmine Sonuçları", redmine_count)
                    
                    with col3:
                        mattermost_count = len(df[df['Source_Platform'] == 'Mattermost'])
                        st.metric("Mattermost Sonuçları", mattermost_count)
                    
                    # Excel'e kaydetme
                    if not df.empty:
                        st.header("💾 Sonuçları İndir")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            csv = df.to_csv(index=False)
                            st.download_button(
                                label="📊 CSV olarak indir",
                                data=csv,
                                file_name=f"{search_term}_arama_sonuclari.csv",
                                mime="text/csv"
                            )
                        
                        with col2:
                            # Excel için openpyxl gerekli
                            try:
                                output = pd.ExcelWriter('temp.xlsx', engine='openpyxl')
                                df.to_excel(output, index=False)
                                output.close()
                                
                                with open('temp.xlsx', 'rb') as f:
                                    excel_data = f.read()
                                
                                st.download_button(
                                    label="📊 Excel olarak indir",
                                    data=excel_data,
                                    file_name=f"{search_term}_arama_sonuclari.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                )
                                
                                # Geçici dosyayı sil
                                os.remove('temp.xlsx')
                                
                            except Exception as e:
                                st.error(f"Excel dosyası oluşturulamadı: {e}")
                
            except Exception as e:
                st.error(f"Veri çekilirken bir hata oluştu: {e}")
    
    # Eğer önceden arama yapılmışsa ve sonuçlar varsa, filtreleme uygula
    elif st.session_state.search_results is not None:
        df = st.session_state.search_results
        
        if not df.empty:
            st.success(f"'{st.session_state.search_term}' için {len(df)} sonuç bulundu!")
            
            # Filtreleme uygula
            filtered_df = df.copy()
            
            # Platform filtresi uygula (eğer değişmişse)
            if platform_filter != "Tümü":
                filtered_df = filtered_df[filtered_df['Source_Platform'] == platform_filter]
            
            if text_filter:
                mask = (filtered_df['Title'].str.contains(text_filter, case=False, na=False) | 
                       filtered_df['Description'].str.contains(text_filter, case=False, na=False))
                filtered_df = filtered_df[mask]
            
            # Mattermost durum filtresi uygula
            if mattermost_status_filter != "Tümü":
                filtered_df = filter_mattermost_by_status(filtered_df, mattermost_status_filter)
            
            # Tarih filtresi uygula
            if date_filter_enabled and (start_date is not None or end_date is not None):
                # Creation_Date sütununu datetime'a çevir
                filtered_df['Creation_Date_dt'] = pd.to_datetime(filtered_df['Creation_Date'], errors='coerce')
                
                if start_date is not None:
                    start_date_dt = pd.to_datetime(start_date)
                    filtered_df = filtered_df[filtered_df['Creation_Date_dt'] >= start_date_dt]
                
                if end_date is not None:
                    end_date_dt = pd.to_datetime(end_date) + pd.Timedelta(days=1)  # Bitiş gününü dahil et
                    filtered_df = filtered_df[filtered_df['Creation_Date_dt'] < end_date_dt]
                
                # Geçici sütunu sil
                filtered_df = filtered_df.drop('Creation_Date_dt', axis=1)
            
            # Sonuçları göster
            st.header("📊 Bulunan Sonuçlar")
            st.subheader(f"Filtrelenmiş Sonuçlar ({len(filtered_df)} adet)")
            
            if not filtered_df.empty:
                # Sonuçları tarihe göre sırala
                filtered_df = filtered_df.sort_values(by='Creation_Date', ascending=False)
                
                for index, row in filtered_df.iterrows():
                    with st.container():
                        # Platform badge'i
                        platform_badge = "🔴 Redmine" if row['Source_Platform'] == 'Redmine' else "💬 Mattermost"
                        
                        st.markdown(f"### {platform_badge} - {row.get('Title', 'Başlık Yok')}")
                        
                        with st.expander("📋 Detayları Gör", expanded=False):
                            col1, col2 = st.columns([2, 1])
                            
                            with col1:
                                st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                                st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                                st.markdown(f"**🏷️ Tip:** {row.get('Content_Type', 'N/A')}")
                                
                                if row.get('Source_Platform') == 'Redmine':
                                    st.markdown(f"**🆔 ID:** {row.get('ID', 'N/A')}")
                                else:
                                    st.markdown(f"**📺 Kanal ID:** {row.get('Channel_ID', 'N/A')}")
                            
                            with col2:
                                st.markdown("**📊 İstatistikler:**")
                                description_length = len(str(row.get('Description', '')))
                                st.markdown(f"• Açıklama uzunluğu: {description_length} karakter")
                            
                            st.markdown("---")
                            st.markdown("**📝 Açıklama/Mesajlar:**")
                            st.text_area("", value=row.get('Description', ''), height=200, disabled=True)
                        
                        st.markdown("---")
    
    # Kullanım talimatları
    else:
        st.info("👆 Yukarıdaki alana aramak istediğiniz iş ismini yazın ve 'Ara' butonuna tıklayın.")

def chatbot_search_interface():
    """Chatbot arama arayüzü"""
    st.header("💬 Chatbot Sorgusu")
    
    # Sorgu girişi
    col1, col2 = st.columns([4, 1])
    
    with col1:
        user_query = st.text_area(
            "Sorgunuzu yazın:",
            placeholder="Örnek: Bana #atp adı altında mattermost kanalında tamamlanan işlerin mesajlarını verir misin?",
            height=100
        )
    
    with col2:
        st.write("")
        st.write("")
        search_button = st.button("🔍 Ara", type="primary", key="chatbot_search")
    
    # Sorgu işleme
    if search_button and user_query:
        with st.spinner("Sorgu analiz ediliyor ve veriler aranıyor..."):
            # Sorguyu parse et
            query_info = parse_natural_language_query(user_query)
            
            if not query_info['search_term']:
                st.error("❌ Sorgunuzda hashtag bulunamadı! Lütfen # ile başlayan bir terim ekleyin.")
                return
            
            # Veri çek
            df = get_all_data_fast(query_info['search_term'], query_info['platform_filter'])
            
            if df.empty:
                st.warning(f"'{query_info['search_term']}' için hiçbir sonuç bulunamadı.")
                return
            
            # Filtreleme uygula
            filtered_df = df.copy()
            
            # Durum filtresi
            if query_info['status_filter'] != "Tümü":
                if query_info['status_filter'] == "Tamamlanan":
                    filtered_df = filtered_df[filtered_df['Status'] == 'tamamlandi']
                elif query_info['status_filter'] == "Devam Eden":
                    filtered_df = filtered_df[filtered_df['Status'] == 'devam_ediyor']
            
            # Tarih filtresi
            if query_info['date_filter']:
                filtered_df['Creation_Date_dt'] = pd.to_datetime(filtered_df['Creation_Date'], errors='coerce')
                filtered_df = filtered_df[filtered_df['Creation_Date_dt'].dt.date >= query_info['date_filter']]
                filtered_df = filtered_df.drop('Creation_Date_dt', axis=1)
            
            # Sonuçları sakla
            st.session_state.last_results = filtered_df
            st.session_state.last_query_info = query_info
            
            # Chat history'ye ekle
            st.session_state.chat_history.append({
                'query': user_query,
                'results_count': len(filtered_df),
                'timestamp': datetime.datetime.now()
            })
            
            # Sonuçları göster
            st.success(f"✅ {len(filtered_df)} sonuç bulundu!")
            
            # Özet bilgiler
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Toplam Sonuç", len(filtered_df))
            
            with col2:
                redmine_count = len(filtered_df[filtered_df['Source_Platform'] == 'Redmine'])
                st.metric("Redmine", redmine_count)
            
            with col3:
                mattermost_count = len(filtered_df[filtered_df['Source_Platform'] == 'Mattermost'])
                st.metric("Mattermost", mattermost_count)
            
            with col4:
                if mattermost_count > 0:
                    mattermost_data = filtered_df[filtered_df['Source_Platform'] == 'Mattermost']
                    tamamlanan_count = len(mattermost_data[mattermost_data['Status'] == 'tamamlandi'])
                    st.metric("Tamamlanan", tamamlanan_count)
                else:
                    st.metric("Tamamlanan", 0)
    
    # Sonuçları göster
    if st.session_state.last_results is not None:
        st.header("📊 Bulunan Sonuçlar")
        
        # Word dokümanı oluştur ve indir
        if st.button("📄 Word Dokümanı Oluştur"):
            with st.spinner("Word dokümanı oluşturuluyor..."):
                doc = create_word_document(st.session_state.last_results, st.session_state.last_query_info)
                
                # Dosyayı kaydet
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                # İndirme butonu
                st.download_button(
                    label="📄 Word Dokümanını İndir",
                    data=docx_buffer.getvalue(),
                    file_name=f"is_takip_raporu_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Sonuçları tablo olarak göster
        st.subheader("Detaylı Sonuçlar")
        
        # Mattermost ve Redmine ayrı ayrı göster
        mattermost_data = st.session_state.last_results[st.session_state.last_results['Source_Platform'] == 'Mattermost']
        redmine_data = st.session_state.last_results[st.session_state.last_results['Source_Platform'] == 'Redmine']
        
        if not mattermost_data.empty:
            st.markdown("**💬 Mattermost Sonuçları:**")
            for index, row in mattermost_data.iterrows():
                status_emoji = "🟢" if row['Status'] == 'tamamlandi' else "🔵"
                with st.expander(f"{status_emoji} {row.get('Title', 'Başlık Yok')}", expanded=False):
                    st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                    st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                    st.markdown(f"**📝 Mesajlar:**")
                    st.text_area("", value=row.get('Description', ''), height=150, disabled=True, key=f"chatbot_mattermost_{index}")
        
        if not redmine_data.empty:
            st.markdown("**🔴 Redmine Sonuçları:**")
            for index, row in redmine_data.iterrows():
                with st.expander(f"🔴 {row.get('Title', 'Başlık Yok')}", expanded=False):
                    st.markdown(f"**📅 Tarih:** {row.get('Creation_Date', 'N/A')}")
                    st.markdown(f"**👤 Yazar:** {row.get('Author', 'N/A')}")
                    st.markdown(f"**📝 Açıklama:**")
                    st.text_area("", value=row.get('Description', ''), height=150, disabled=True, key=f"chatbot_redmine_{index}")
    
    # Chat history
    if st.session_state.chat_history:
        st.header("💬 Sorgu Geçmişi")
        for i, chat in enumerate(reversed(st.session_state.chat_history)):
            with st.expander(f"🔍 {chat['query'][:50]}... ({chat['results_count']} sonuç)", expanded=False):
                st.markdown(f"**Sorgu:** {chat['query']}")
                st.markdown(f"**Sonuç Sayısı:** {chat['results_count']}")
                st.markdown(f"**Tarih:** {chat['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}")

# ==============================================================================
# 6. MATTERMOST İŞ DURUMU ANALİZ FONKSİYONU
# ==============================================================================
def analyze_mattermost_status(thread_data):
    """
    Mattermost thread'inin durumunu belirler
    Artık durum bilgisi DataFrame'de Status sütununda var
    """
    if thread_data is None or 'Status' not in thread_data:
        return "devam_ediyor"
    
    status = thread_data['Status']
    if pd.isna(status) or status == "":
        return "devam_ediyor"
    
    return status

# ==============================================================================
# 7. MATTERMOST DURUM FİLTRELEME FONKSİYONU
# ==============================================================================
def filter_mattermost_by_status(df, status_filter):
    """
    Mattermost sonuçlarını durumlarına göre filtreler
    """
    if status_filter == "Tümü":
        return df
    
    # Sadece Mattermost sonuçlarını al
    mattermost_df = df[df['Source_Platform'] == 'Mattermost'].copy()
    
    if mattermost_df.empty:
        return df
    
    # Status sütunu zaten fast_mattermost_fetch'te set edildi, tekrar hesaplamaya gerek yok
    
    # Durum filtresini uygula
    if status_filter == "Tamamlanan":
        filtered_df = mattermost_df[mattermost_df['Status'] == 'tamamlandi']
    elif status_filter == "Devam Eden":
        filtered_df = mattermost_df[mattermost_df['Status'] == 'devam_ediyor']
    else:
        filtered_df = mattermost_df
    
    # Status sütununu kaldır
    if 'Status' in filtered_df.columns:
        filtered_df = filtered_df.drop('Status', axis=1)
    
    # Redmine sonuçlarını da ekle
    redmine_df = df[df['Source_Platform'] == 'Redmine']
    final_df = pd.concat([redmine_df, filtered_df], ignore_index=True)
    
    return final_df

if __name__ == "__main__":
    main() 